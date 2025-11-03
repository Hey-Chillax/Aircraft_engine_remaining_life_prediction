"""
读取 Word 文档内容的脚本
"""
from docx import Document
import sys

def read_docx(file_path):
    """读取 Word 文档并提取所有文本内容"""
    try:
        doc = Document(file_path)
        
        print("=" * 80)
        print(f"文档路径: {file_path}")
        print("=" * 80)
        print()
        
        # 提取所有段落
        full_text = []
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:  # 只打印非空段落
                full_text.append(text)
                print(f"[段落 {i}]")
                print(text)
                print()
        
        print("=" * 80)
        print(f"总段落数: {len(doc.paragraphs)}")
        print(f"非空段落数: {len(full_text)}")
        print("=" * 80)
        
        # 提取表格内容
        if doc.tables:
            print("\n" + "=" * 80)
            print("表格内容:")
            print("=" * 80)
            for i, table in enumerate(doc.tables, 1):
                print(f"\n[表格 {i}]")
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    print(" | ".join(row_data))
                print()
        
        return full_text
        
    except Exception as e:
        print(f"读取文档时出错: {e}")
        import traceback
        traceback.print_exc()
        return None

if __name__ == "__main__":
    file_path = "documentation/航空发动机剩余寿命预测2025.docx"
    read_docx(file_path)

